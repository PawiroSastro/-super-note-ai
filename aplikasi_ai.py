# ======================================================
# AI SUPER NOTE PRO
# ======================================================

import streamlit as st
import os
from openai import OpenAI
import streamlit_mermaid as st_mermaid
from pytubefix import YouTube
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

# ---------------- CONFIG ----------------

st.set_page_config(
    page_title="AI Super Note",
    page_icon="🚀",
    layout="wide"
)

MY_SECRET_KEY = st.secrets["OPENAI_API_KEY"]

client = OpenAI(api_key=MY_SECRET_KEY)

# ---------------- STYLE ----------------

st.markdown("""
<style>

#MainMenu {visibility:hidden;}
footer {visibility:hidden;}
header {visibility:hidden;}

[data-testid="stToolbar"]{
visibility:hidden;
}

[data-testid="stDecoration"]{
visibility:hidden;
}

[data-testid="stStatusWidget"]{
visibility:hidden;
}

.stApp{
background:#f4f7fb;
}

.block-container{
padding-top:1rem;
max-width:1100px;
}

.hero{
background:linear-gradient(135deg,#0f4c81,#2563eb);
padding:22px;
border-radius:25px;
color:white;
margin-bottom:20px;
box-shadow:0 8px 30px rgba(0,0,0,.2);
}

.hero h1{
margin:0;
font-size:42px;
line-height:1.1;
}

.hero p{
font-size:18px;
opacity:.95;
margin-top:10px;
}

/* MODE HP */
@media (max-width:768px){

.hero{
padding:18px;
border-radius:20px;
}

.hero h1{
font-size:26px;
}

.hero p{
font-size:14px;
}

.block-container{
padding-top:.5rem;
padding-left:1rem;
padding-right:1rem;
}

}

.stButton>button{
width:100%;
height:55px;
border-radius:15px;
font-size:18px;
font-weight:bold;
background:#0f4c81;
color:white;
border:none;
}

.stDownloadButton>button{
width:100%;
border-radius:15px;
background:#16a34a;
color:white;
font-weight:bold;
border:none;
}

</style>
""", unsafe_allow_html=True)

# ---------------- HEADER ----------------

st.markdown("""
<div class='hero'>
<h1>🚀 AI Super Note Pro</h1>
<p>
Transkripsi • Ringkasan AI • Mindmap • PDF • Word
</p>
</div>
""", unsafe_allow_html=True)

# ---------------- SESSION ----------------

if "data" not in st.session_state:

    st.session_state.data={

    "audio":None,
    "transkrip":"",
    "ringkasan":"",
    "mindmap":""

    }

# ---------------- INPUT ----------------

c1,c2,c3=st.columns(3)

with c1:

    st.subheader("🎤 Rekam")

    rec=st.audio_input(
    "Gunakan mikrofon"
    )

    if rec:

        st.session_state.data["audio"]=rec.read()


with c2:

    st.subheader("📂 Upload")

    up=st.file_uploader(
    "Upload Audio",
    type=["mp3","wav","ogg","m4a"]
    )

    if up:

        st.session_state.data["audio"]=up.read()

with c3:

    st.subheader("🔗 YouTube")

    yt=st.text_input(
    "Tempel link"
    )

# ---------------- PROSES ----------------

st.divider()

if st.button(
"🚀 LUNCURKAN SASARAN KE TARGET"
):

    with st.spinner("Memproses AI..."):

        try:

            if yt:

                y=YouTube(yt)

                y.streams.get_audio_only().download(
                filename="yt.mp3"
                )

                with open(
                "yt.mp3",
                "rb"
                ) as f:

                    st.session_state.data["audio"]=f.read()

                os.remove(
                "yt.mp3"
                )

            if st.session_state.data["audio"]:

                with open(
                "temp.wav",
                "wb"
                ) as f:

                    f.write(
                    st.session_state.data["audio"]
                    )

                with open(
                "temp.wav",
                "rb"
                ) as audio:

                    hasil=client.audio.transcriptions.create(
                    model="gpt-4o-mini-transcribe",
                    file=audio,
                    language="id"
                    )

                teks=hasil.text

                st.session_state.data["transkrip"]=teks

                ai=client.chat.completions.create(

                model="gpt-4o-mini",

                messages=[

                {

                "role":"user",

                "content":
                f"""
Ringkas:

{teks}

buat:
1 ringkasan
2 poin penting
3 insight
"""
                }

                ]

                )

                st.session_state.data["ringkasan"]=ai.choices[0].message.content

                st.session_state.data["mindmap"]="""
graph TD
A[AI Super Note]
-->B[Transkrip]
-->C[Ringkasan]
-->D[Mindmap]
"""

                st.rerun()

        except Exception as e:

            st.error(e)

# ---------------- OUTPUT ----------------

d=st.session_state.data

if d["transkrip"]:

    a,b=st.columns([2,1])

    with a:

        st.subheader("🧠 Visual")

        st_mermaid.st_mermaid(
        d["mindmap"]
        )

        st.info(
        d["ringkasan"]
        )

        st.text_area(
        "Transkrip",
        d["transkrip"],
        height=250
        )

    with b:

        st.subheader(
        "💾 Export"
        )

        doc=Document()

        doc.add_heading(
        "AI SUPER NOTE",
        level=1
        )

        doc.add_paragraph(
        d["transkrip"]
        )

        bio=BytesIO()

        doc.save(
        bio
        )

        st.download_button(
        "📝 Word",
        bio.getvalue(),
        "laporan.docx"
        )

        p=BytesIO()

        pdf=SimpleDocTemplate(
        p,
        pagesize=A4
        )

        style=getSampleStyleSheet()

        pdf.build([

        Paragraph(
        "Laporan AI",
        style["Title"]
        ),

        Paragraph(
        d["transkrip"],
        style["Normal"]
        )

        ])

        st.download_button(
        "📕 PDF",
        p.getvalue(),
        "laporan.pdf"
        )
